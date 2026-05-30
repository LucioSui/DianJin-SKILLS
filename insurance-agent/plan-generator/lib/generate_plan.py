"""保险计划书 DOCX 生成引擎。

使用方式:
    # 独立输入
    python generate_plan.py --input plan_input.json --output 计划书.docx

    # 技能联动
    python generate_plan.py --profile 画像.json --recommendation 推荐.json --agent 代理人.json --output 计划书.docx
"""

import json
import argparse
import os
import sys
import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# 本项目内部导入
from .benefit_calculator import calculate_benefit, _load_models

# ============================================================
# 模板/配置路径
# ============================================================
_TEMPLATES_DIR = os.path.join(_SCRIPT_DIR, "..", "templates")
_COMPLIANCE_PATH = os.path.join(_TEMPLATES_DIR, "compliance_text.json")

# ============================================================
# 颜色常量
# ============================================================
BLUE_PRIMARY = (46, 117, 182)   # #2E75B6
BLUE_DARK = (31, 78, 121)      # #1F4E79
GRAY_TEXT = (100, 100, 100)
GRAY_LIGHT = (160, 160, 160)

# ============================================================
# DOCX 辅助函数 (复用自 gen_skill_intro.py)
# ============================================================

def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def set_run_style(run, font_name='微软雅黑', size=None, bold=False, color=None):
    """设置文字样式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h


def add_para(doc, text, font_size=11, bold=False, color=None, alignment=None, space_after=6):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_style(run, size=font_size, bold=bold, color=color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """添加表格 — 蓝色表头 + 交替行色"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_style(run, size=10, bold=True, color=(255, 255, 255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E75B6')

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_style(run, size=10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F2F7FB')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_kv_table(doc, items, col_widths=None):
    """添加键值对表格(2列, 无表头行, 交替色)"""
    table = doc.add_table(rows=len(items), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, (k, v) in enumerate(items):
        ck = table.rows[r_idx].cells[0]
        cv = table.rows[r_idx].cells[1]
        ck.text = ''
        cv.text = ''
        run_k = ck.paragraphs[0].add_run(str(k))
        set_run_style(run_k, size=10, bold=True, color=BLUE_DARK)
        set_cell_shading(ck, 'E8F0FE')
        run_v = cv.paragraphs[0].add_run(str(v))
        set_run_style(run_v, size=10)
        if r_idx % 2 == 1:
            set_cell_shading(cv, 'F2F7FB')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def fmt(v):
    """格式化金额"""
    if isinstance(v, float):
        if v == int(v):
            return f"{int(v):,}"
        return f"{v:,.1f}"
    if isinstance(v, int):
        return f"{v:,}"
    return str(v)


# ============================================================
# 保障缺口计算 (回退逻辑, 复用 analyze_profile.py 公式)
# ============================================================

def calculate_gap(customer):
    """计算五维保障缺口（当输入未提供 gap_analysis 时使用）"""
    income = customer.get("annual_income", 0)
    policies = customer.get("existing_policies", [])
    age = customer.get("age", 0)
    monthly_expense = customer.get("monthly_expense", 0)

    def get_sum_by_type(t):
        return sum(p.get("sum_assured", 0) for p in policies if t in p.get("type", ""))

    life_suggested = income * 10
    life_existing = get_sum_by_type("寿险")
    life_gap = max(0, life_suggested - life_existing)

    ci_suggested = income * 5 + 300000
    ci_existing = get_sum_by_type("重疾")
    ci_gap = max(0, ci_suggested - ci_existing)

    med_suggested = 1000000
    med_existing = get_sum_by_type("医疗")
    med_gap = max(0, med_suggested - med_existing)
    med_covered = med_existing >= med_suggested

    acc_suggested = income * 10
    acc_existing = get_sum_by_type("意外")
    acc_gap = max(0, acc_suggested - acc_existing)

    retire_years = max(0, 80 - max(age, 60))
    pension_suggested = monthly_expense * 12 * retire_years
    pension_existing = get_sum_by_type("年金")
    pension_gap = max(0, pension_suggested - pension_existing)

    def gap_rate(suggested, existing):
        if suggested == 0:
            return 0
        return round((max(0, suggested - existing) / suggested) * 100, 1)

    def urgency(rate):
        if rate >= 80:
            return "紧急"
        elif rate >= 50:
            return "重要"
        return "优化"

    result = {
        "寿险": {"建议保额": round(life_suggested / 10000, 1), "已有保额": round(life_existing / 10000, 1),
                 "缺口": round(life_gap / 10000, 1), "缺口率": gap_rate(life_suggested, life_existing),
                 "紧迫度": urgency(gap_rate(life_suggested, life_existing))},
        "重疾险": {"建议保额": round(ci_suggested / 10000, 1), "已有保额": round(ci_existing / 10000, 1),
                   "缺口": round(ci_gap / 10000, 1), "缺口率": gap_rate(ci_suggested, ci_existing),
                   "紧迫度": urgency(gap_rate(ci_suggested, ci_existing))},
        "医疗险": {"建议保额": round(med_suggested / 10000, 1), "已有保额": round(med_existing / 10000, 1),
                   "缺口": round(med_gap / 10000, 1), "缺口率": gap_rate(med_suggested, med_existing),
                   "紧迫度": "优化" if med_covered else urgency(gap_rate(med_suggested, med_existing)),
                   "已覆盖": med_covered},
        "意外险": {"建议保额": round(acc_suggested / 10000, 1), "已有保额": round(acc_existing / 10000, 1),
                   "缺口": round(acc_gap / 10000, 1), "缺口率": gap_rate(acc_suggested, acc_existing),
                   "紧迫度": urgency(gap_rate(acc_suggested, acc_existing))},
        "养老储备": {"建议保额": round(pension_suggested / 10000, 1), "已有保额": round(pension_existing / 10000, 1),
                    "缺口": round(pension_gap / 10000, 1), "缺口率": gap_rate(pension_suggested, pension_existing),
                    "紧迫度": urgency(gap_rate(pension_suggested, pension_existing))},
    }

    total_s = life_suggested + ci_suggested + med_suggested + acc_suggested + pension_suggested
    total_e = life_existing + ci_existing + med_existing + acc_existing + pension_existing
    result["综合缺口率"] = gap_rate(total_s, total_e)
    return result


# ============================================================
# 章节生成函数
# ============================================================

def add_cover_page(doc, data):
    """第1章: 封面"""
    for _ in range(4):
        doc.add_paragraph()

    agent = data.get("agent", {})
    company = agent.get("company", "")
    if company:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company)
        set_run_style(run, size=18, bold=True, color=BLUE_DARK)

    doc.add_paragraph()

    customer = data.get("customer", {})
    name = customer.get("name", "客户")
    gender = customer.get("gender", "")
    honorific = "先生" if gender == "男" else ("女士" if gender == "女" else "")
    config = data.get("plan_config", {})
    title = config.get("plan_title", f"{name}{honorific} 家庭保障规划方案")

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run(title)
    set_run_style(run, size=26, bold=True, color=BLUE_PRIMARY)

    doc.add_paragraph()

    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_p.add_run('_' * 50)
    set_run_style(run, size=11, color=BLUE_PRIMARY)

    doc.add_paragraph()

    plan_date = config.get("plan_date", datetime.date.today().strftime("%Y-%m-%d"))
    info_lines = [
        f"日期: {plan_date}",
        f"服务顾问: {agent.get('name', '')}",
        "本计划书为客户专属定制，请妥善保管",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_style(run, size=12, color=GRAY_TEXT)

    doc.add_page_break()


def add_toc(doc, sections):
    """第2章: 目录"""
    add_heading_styled(doc, '目录', level=1)

    section_names = {
        "needs": "客户需求分析",
        "gap": "保障缺口分析",
        "products": "产品方案推荐",
        "benefits": "利益演示",
        "budget": "保费预算汇总",
        "compliance": "合规声明与重要提示",
        "agent": "您的专属服务顾问",
    }

    num = 1
    for key in ["needs", "gap", "products", "benefits", "budget", "compliance", "agent"]:
        if key in sections:
            label = section_names.get(key, key)
            add_para(doc, f"{'一二三四五六七八九十'[num-1]}、{label}", font_size=12, space_after=8)
            num += 1

    doc.add_page_break()


def add_needs_analysis(doc, data):
    """第3章: 客户需求分析"""
    add_heading_styled(doc, '客户需求分析', level=1)

    customer = data.get("customer", {})
    narratives = data.get("narratives", {})

    # 叙述段落
    if narratives.get("needs_narrative"):
        add_para(doc, narratives["needs_narrative"], font_size=11, space_after=10)

    # 3.1 基本信息
    add_heading_styled(doc, '基本信息', level=2)
    gender = customer.get("gender", "")
    honorific = "先生" if gender == "男" else ("女士" if gender == "女" else "")
    items = [
        ("姓名", f'{customer.get("name", "")}{honorific}'),
        ("性别", gender),
        ("年龄", f'{customer.get("age", "")}岁'),
        ("职业", customer.get("occupation", "")),
        ("所在城市", customer.get("city", "")),
        ("婚姻状况", customer.get("marital_status", "")),
        ("生命周期阶段", customer.get("life_stage", "")),
        ("风险偏好", customer.get("risk_preference", "")),
    ]
    items = [(k, v) for k, v in items if v and str(v) != "岁"]
    add_kv_table(doc, items, col_widths=[4, 11])

    # 3.2 家庭成员
    family = customer.get("family_members", [])
    if family:
        add_heading_styled(doc, '家庭成员', level=2)
        add_table(doc,
                  ["关系", "姓名", "年龄", "职业"],
                  [[m.get("relation", ""), m.get("name", ""), f'{m.get("age", "")}岁', m.get("occupation", "")]
                   for m in family],
                  col_widths=[3, 3.5, 3, 5.5])

    # 3.3 家庭财务概况
    income = customer.get("annual_income")
    expense = customer.get("monthly_expense")
    if income or expense:
        add_heading_styled(doc, '家庭财务概况', level=2)
        fin_items = []
        if income:
            fin_items.append(("家庭年收入", f"{fmt(income)}元"))
        if expense:
            fin_items.append(("月均支出", f"{fmt(expense)}元"))
            fin_items.append(("年支出", f"{fmt(expense * 12)}元"))
        assets = customer.get("assets", {})
        if assets.get("mortgage_remaining"):
            fin_items.append(("房贷余额", f"{fmt(assets['mortgage_remaining'])}元"))
        if assets.get("liquid_assets"):
            fin_items.append(("流动资产", f"{fmt(assets['liquid_assets'])}元"))
        add_kv_table(doc, fin_items, col_widths=[4, 11])

    # 3.4 现有保障盘点
    policies = customer.get("existing_policies", [])
    if policies:
        add_heading_styled(doc, '现有保障盘点', level=2)
        rows = []
        total_prem = 0
        for p in policies:
            prem = p.get("annual_premium", 0)
            total_prem += prem
            rows.append([
                p.get("type", ""),
                p.get("product_name", ""),
                f'{fmt(p.get("sum_assured", 0))}元',
                f'{fmt(prem)}元/年',
                p.get("status", ""),
            ])
        add_table(doc,
                  ["险种", "产品名称", "保额", "年缴保费", "状态"],
                  rows, col_widths=[2.5, 4, 3, 3, 2.5])

        summary = f"现有保障年缴保费合计: {fmt(total_prem)}元"
        if income:
            ratio = total_prem / income * 100
            summary += f"，占家庭年收入的 {ratio:.1f}%"
        add_para(doc, summary, font_size=10, color=GRAY_TEXT)


def add_gap_analysis(doc, data):
    """第4章: 保障缺口分析"""
    add_heading_styled(doc, '保障缺口分析', level=1)

    gap = data.get("gap_analysis")
    customer = data.get("customer", {})
    if not gap:
        gap = calculate_gap(customer)

    narratives = data.get("narratives", {})
    if narratives.get("gap_narrative"):
        add_para(doc, narratives["gap_narrative"], font_size=11, space_after=10)

    # 4.1 五维缺口评分总览
    add_heading_styled(doc, '五维缺口评分总览', level=2)

    dims = ["寿险", "重疾险", "医疗险", "意外险", "养老储备"]
    headers = ["维度", "建议保额(万)", "已有保额(万)", "缺口(万)", "缺口率", "紧迫度"]
    rows = []
    for d in dims:
        g = gap.get(d, {})
        urgency = g.get("紧迫度", "")
        urgency_label = f"[{urgency}]" if urgency else ""
        rows.append([
            d,
            fmt(g.get("建议保额", 0)),
            fmt(g.get("已有保额", 0)),
            fmt(g.get("缺口", 0)),
            f'{g.get("缺口率", 0)}%',
            urgency_label,
        ])
    add_table(doc, headers, rows, col_widths=[2.5, 2.5, 2.5, 2.5, 2, 2.5])

    overall = gap.get("综合缺口率", 0)
    if overall:
        add_para(doc, f"综合缺口率: {overall}%", font_size=11, bold=True, color=BLUE_PRIMARY)

    # 4.2 可视化(文本柱状图)
    add_heading_styled(doc, '缺口分布', level=2)
    max_rate = max(gap.get(d, {}).get("缺口率", 0) for d in dims)
    max_rate = max(max_rate, 1)
    for d in dims:
        g = gap.get(d, {})
        rate = g.get("缺口率", 0)
        bar_len = int(rate / max_rate * 20)
        bar = "█" * bar_len + "░" * (20 - bar_len)
        add_para(doc, f"{d:　<4} {bar} {rate}%", font_size=10, space_after=2)

    # 4.3 风险优先级排序
    add_heading_styled(doc, '风险优先级排序', level=2)
    priority_order = {"紧急": 1, "高": 1, "重要": 2, "中": 2, "优化": 3, "低": 3}
    sorted_dims = sorted(dims, key=lambda d: priority_order.get(gap.get(d, {}).get("紧迫度", ""), 9))

    for d in sorted_dims:
        g = gap.get(d, {})
        urg = g.get("紧迫度", "")
        gap_val = g.get("缺口", 0)
        if gap_val > 0:
            p = doc.add_paragraph()
            marker = f"[{urg}]" if urg else ""
            run = p.add_run(f"  {marker} {d}: 缺口 {fmt(gap_val)}万元")
            urg_color = (231, 76, 60) if urg in ("紧急", "高") else (BLUE_PRIMARY if urg in ("重要", "中") else GRAY_TEXT)
            set_run_style(run, size=11, color=urg_color)

    doc.add_paragraph()


def add_product_plans(doc, data):
    """第5章: 产品方案推荐"""
    add_heading_styled(doc, '产品方案推荐', level=1)

    plans = data.get("product_plans", {})
    narratives = data.get("narratives", {})

    if narratives.get("recommendation_narrative"):
        add_para(doc, narratives["recommendation_narrative"], font_size=11, space_after=10)

    for plan_key in ["plan_a", "plan_b"]:
        plan = plans.get(plan_key)
        if not plan:
            continue

        label = plan.get("plan_label", "方案A" if plan_key == "plan_a" else "方案B")
        add_heading_styled(doc, label, level=2)

        products = plan.get("products", [])
        if not products:
            add_para(doc, "（暂无产品配置）", color=GRAY_TEXT)
            continue

        headers = ["产品名称", "险种", "被保人", "保额", "年缴保费", "缴费期", "保障期"]
        rows = []
        total_premium = 0
        for p in products:
            prem = p.get("annual_premium", 0)
            total_premium += prem
            rows.append([
                p.get("product_name", ""),
                _type_label(p.get("product_type", "")),
                p.get("insured_person", ""),
                f'{fmt(p.get("sum_assured", 0))}元',
                f'{fmt(prem)}元',
                p.get("payment_period", ""),
                p.get("coverage_period", ""),
            ])

        add_table(doc, headers, rows, col_widths=[3, 2, 2, 2.5, 2.5, 1.5, 1.5])

        add_para(doc, f"年缴保费合计: {fmt(total_premium)}元", font_size=11, bold=True, color=BLUE_PRIMARY)

    # 双方案对比
    if plans.get("plan_a") and plans.get("plan_b"):
        add_heading_styled(doc, '方案对比', level=2)
        pa = plans["plan_a"]
        pb = plans["plan_b"]
        prem_a = sum(p.get("annual_premium", 0) for p in pa.get("products", []))
        prem_b = sum(p.get("annual_premium", 0) for p in pb.get("products", []))
        cnt_a = len(pa.get("products", []))
        cnt_b = len(pb.get("products", []))

        add_table(doc,
                  ["对比项", pa.get("plan_label", "方案A"), pb.get("plan_label", "方案B")],
                  [
                      ["产品数量", f"{cnt_a}款", f"{cnt_b}款"],
                      ["年缴保费", f"{fmt(prem_a)}元", f"{fmt(prem_b)}元"],
                      ["定位", "基础保障", "全面保障"],
                  ],
                  col_widths=[3.5, 5.5, 5.5])


def add_benefit_demo(doc, data):
    """第6章: 利益演示"""
    add_heading_styled(doc, '利益演示', level=1)

    add_para(doc, "以下利益演示基于假设条件测算，不代表确定收益。实际保障以保险合同条款为准。",
             font_size=10, color=GRAY_TEXT, space_after=10)

    plans = data.get("product_plans", {})
    customer = data.get("customer", {})
    insured_age = customer.get("age", 30)
    models = _load_models()

    section_num = 1
    seen_types = set()

    for plan_key in ["plan_a", "plan_b"]:
        plan = plans.get(plan_key)
        if not plan:
            continue
        for product in plan.get("products", []):
            pt = product.get("product_type", "")
            pname = product.get("product_name", pt)

            # 避免同类型重复 (如方案A和B都有定寿)
            type_key = f"{pt}_{pname}"
            if type_key in seen_types:
                continue
            seen_types.add(type_key)

            # 计算利益
            bd = product.get("benefit_details", {})
            # 确定被保人年龄
            ip_age = insured_age
            ip = product.get("insured_person", "")
            if ip:
                for fm in customer.get("family_members", []):
                    if fm.get("name") == ip or fm.get("relation") == ip:
                        ip_age = fm.get("age", insured_age)
                        break

            payment_str = product.get("payment_period", "20年")
            pay_years = _parse_years(payment_str)
            cov_str = product.get("coverage_period", "终身")
            cov_years = 999 if "终身" in cov_str else _parse_years(cov_str)

            result = calculate_benefit(
                product_type=pt,
                benefit_details=bd,
                product_name=pname,
                annual_premium=product.get("annual_premium", 0),
                insured_age=ip_age,
                payment_years=pay_years,
                coverage_years=cov_years,
                sum_assured=product.get("sum_assured", 0),
                models=models,
            )

            add_heading_styled(doc, f'{section_num}. {pname} 利益演示', level=2)
            add_table(doc, result["headers"], result["rows"])

            # 注释
            for note in result.get("notes", []):
                add_para(doc, note, font_size=9, color=GRAY_LIGHT, space_after=2)

            section_num += 1


def add_budget_summary(doc, data):
    """第7章: 保费预算汇总"""
    add_heading_styled(doc, '保费预算汇总', level=1)

    plans = data.get("product_plans", {})
    customer = data.get("customer", {})
    income = customer.get("annual_income", 0)

    # 7.1 保费明细
    add_heading_styled(doc, '保费明细', level=2)

    all_products = []
    for plan_key in ["plan_a", "plan_b"]:
        plan = plans.get(plan_key)
        if not plan:
            continue
        label = plan.get("plan_label", plan_key)
        for p in plan.get("products", []):
            all_products.append((label, p))

    if not all_products:
        add_para(doc, "（暂无产品配置）", color=GRAY_TEXT)
        return

    headers = ["所属方案", "产品名称", "年缴保费", "缴费期"]
    rows = []
    plan_totals = {}
    for label, p in all_products:
        prem = p.get("annual_premium", 0)
        plan_totals[label] = plan_totals.get(label, 0) + prem
        rows.append([label, p.get("product_name", ""), f"{fmt(prem)}元", p.get("payment_period", "")])
    add_table(doc, headers, rows, col_widths=[3, 5, 3.5, 3.5])

    # 7.2 预算分析
    add_heading_styled(doc, '预算分析', level=2)

    for label, total in plan_totals.items():
        items = [
            ("方案", label),
            ("年缴保费合计", f"{fmt(total)}元"),
            ("月均折算", f"{fmt(int(total / 12))}元"),
        ]
        if income > 0:
            ratio = total / income * 100
            if ratio < 10:
                benchmark = "理想 (行业建议 < 10%)"
            elif ratio < 15:
                benchmark = "适中 (行业建议 10-15%)"
            else:
                benchmark = "偏高 (超过15%，建议优化)"
            items.append(("占家庭年收入", f"{ratio:.1f}% — {benchmark}"))
        add_kv_table(doc, items, col_widths=[4, 11])

    # 7.3 缴费时间线
    add_heading_styled(doc, '缴费时间线', level=2)
    today_year = datetime.date.today().year
    headers_tl = ["产品名称", "缴费年限", "预计缴费完成年份"]
    rows_tl = []
    for _, p in all_products:
        pay_str = p.get("payment_period", "")
        pay_years = _parse_years(pay_str)
        end_year = today_year + pay_years if pay_years > 0 else "—"
        rows_tl.append([p.get("product_name", ""), pay_str, str(end_year)])
    add_table(doc, headers_tl, rows_tl, col_widths=[5, 5, 5])


def add_compliance(doc, data):
    """第8章: 合规声明"""
    add_heading_styled(doc, '合规声明与重要提示', level=1)

    # 加载合规文本
    try:
        with open(_COMPLIANCE_PATH, "r", encoding="utf-8") as f:
            texts = json.load(f)
    except FileNotFoundError:
        texts = {"general": ["产品条款以实际合同为准。"], "required_disclaimers": []}

    # 收集产品类型
    product_types = set()
    plans = data.get("product_plans", {})
    for plan_key in ["plan_a", "plan_b"]:
        plan = plans.get(plan_key)
        if not plan:
            continue
        for p in plan.get("products", []):
            product_types.add(p.get("product_type", ""))

    # 通用声明
    for text in texts.get("general", []):
        p = doc.add_paragraph()
        run = p.add_run(f"  {text}")
        set_run_style(run, size=9, color=GRAY_TEXT)
        p.paragraph_format.space_after = Pt(4)

    # 产品类型特定声明
    for pt in sorted(product_types):
        specific = texts.get(pt, [])
        for text in specific:
            p = doc.add_paragraph()
            run = p.add_run(f"  {text}")
            set_run_style(run, size=9, color=GRAY_TEXT)
            p.paragraph_format.space_after = Pt(4)

    # 必选声明
    doc.add_paragraph()
    for text in texts.get("required_disclaimers", []):
        p = doc.add_paragraph()
        run = p.add_run(f"  {text}")
        set_run_style(run, size=9, bold=True, color=GRAY_TEXT)
        p.paragraph_format.space_after = Pt(4)


def add_agent_info(doc, data):
    """第9章: 代理人信息"""
    add_heading_styled(doc, '您的专属服务顾问', level=1)

    agent = data.get("agent", {})
    if not agent:
        add_para(doc, "（未提供代理人信息）", color=GRAY_TEXT)
        return

    items = []
    if agent.get("name"):
        items.append(("姓名", agent["name"]))
    if agent.get("title"):
        items.append(("职位", agent["title"]))
    if agent.get("company"):
        items.append(("所属公司", agent["company"]))
    if agent.get("branch"):
        items.append(("所属机构", agent["branch"]))
    if agent.get("license_number"):
        items.append(("执业证号", agent["license_number"]))
    if agent.get("phone"):
        items.append(("联系电话", agent["phone"]))
    if agent.get("wechat"):
        items.append(("微信", agent["wechat"]))
    if agent.get("email"):
        items.append(("邮箱", agent["email"]))

    if items:
        add_kv_table(doc, items, col_widths=[4, 11])

    add_para(doc, "感谢您的信任。我将竭诚为您和您的家人提供专业、持续的保障服务。如有任何疑问，请随时联系我。",
             font_size=11, color=BLUE_PRIMARY, space_after=10)


def add_end_page(doc):
    """尾页"""
    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()

    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ep.add_run('— End —')
    set_run_style(run, size=14, color=GRAY_LIGHT)

    doc.add_paragraph()
    footer_lines = [
        '本计划书由 Qoder AI 辅助生成',
        f'生成日期: {datetime.date.today().strftime("%Y年%m月%d日")}',
    ]
    for fl in footer_lines:
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(fl)
        set_run_style(run, size=10, color=GRAY_LIGHT)


# ============================================================
# 辅助工具
# ============================================================

_TYPE_LABELS = {
    "term_life": "定期寿险", "whole_life": "终身寿险", "critical_illness": "重疾险",
    "medical": "医疗险", "accident": "意外险", "annuity": "年金险",
    "education_fund": "教育金", "universal_life": "万能险",
}

def _type_label(t):
    return _TYPE_LABELS.get(t, t)

def _parse_years(s):
    """从 '20年' / '至70周岁' 等字符串提取数字"""
    import re
    m = re.search(r'(\d+)', str(s))
    return int(m.group(1)) if m else 20


# ============================================================
# 数据加载与合并
# ============================================================

def load_unified_input(args):
    """根据 CLI 参数加载并合并数据为统一格式"""

    if args.input:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data

    # 联动模式: 合并多个文件
    data = {"plan_config": {}, "customer": {}, "product_plans": {}, "agent": {}}

    if args.profile:
        with open(args.profile, "r", encoding="utf-8") as f:
            profile_data = json.load(f)
        # 支持单客户 dict 或客户列表 (取第一个)
        if isinstance(profile_data, list):
            profile_data = profile_data[0]
        data["customer"] = profile_data

    if args.recommendation:
        with open(args.recommendation, "r", encoding="utf-8") as f:
            rec_data = json.load(f)
        if "product_plans" in rec_data:
            data["product_plans"] = rec_data["product_plans"]
        else:
            data["product_plans"] = rec_data

    if args.agent:
        with open(args.agent, "r", encoding="utf-8") as f:
            data["agent"] = json.load(f)

    if args.narratives:
        with open(args.narratives, "r", encoding="utf-8") as f:
            data["narratives"] = json.load(f)

    return data


# ============================================================
# 主流程
# ============================================================

def generate_plan(data, output_path, sections=None):
    """生成保险计划书 DOCX 文档。

    Args:
        data: 统一输入数据 dict
        output_path: 输出文件路径
        sections: 包含的章节列表，默认全部
    """
    if sections is None:
        sections = ["cover", "needs", "gap", "products", "benefits", "budget", "compliance", "agent"]

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # 确保 gap_analysis 存在
    if "gap_analysis" not in data and "customer" in data:
        customer = data["customer"]
        if customer.get("annual_income") or customer.get("existing_policies"):
            data["gap_analysis"] = calculate_gap(customer)

    # 按章节顺序生成
    if "cover" in sections:
        add_cover_page(doc, data)

    # 目录
    content_sections = [s for s in sections if s not in ("cover",)]
    if content_sections:
        add_toc(doc, content_sections)

    if "needs" in sections:
        add_needs_analysis(doc, data)
        doc.add_page_break()

    if "gap" in sections:
        add_gap_analysis(doc, data)
        doc.add_page_break()

    if "products" in sections:
        add_product_plans(doc, data)
        doc.add_page_break()

    if "benefits" in sections:
        add_benefit_demo(doc, data)
        doc.add_page_break()

    if "budget" in sections:
        add_budget_summary(doc, data)
        doc.add_page_break()

    if "compliance" in sections:
        add_compliance(doc, data)
        doc.add_page_break()

    if "agent" in sections:
        add_agent_info(doc, data)

    add_end_page(doc)

    # 保存
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"计划书已生成: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="保险计划书 DOCX 生成引擎")
    parser.add_argument("--input", help="完整输入 JSON 文件路径")
    parser.add_argument("--profile", help="客户画像 JSON 文件路径(联动模式)")
    parser.add_argument("--recommendation", help="产品推荐 JSON 文件路径(联动模式)")
    parser.add_argument("--agent", help="代理人信息 JSON 文件路径(联动模式)")
    parser.add_argument("--narratives", help="LLM 叙述内容 JSON 文件路径")
    parser.add_argument("--output", default="保障规划方案.docx", help="输出 DOCX 文件路径")
    parser.add_argument("--sections", help="包含的章节(逗号分隔)", default=None)

    args = parser.parse_args()

    if not args.input and not args.profile:
        parser.error("请提供 --input 或 --profile 参数")

    data = load_unified_input(args)

    sections = None
    if args.sections:
        sections = [s.strip() for s in args.sections.split(",")]

    generate_plan(data, args.output, sections)


if __name__ == "__main__":
    main()
